from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = r"D:\Workspace\Arcitecture\操作手册.docx"

NAVY = "1F4E79"
NAVY_DARK = "123852"
BLUE_LIGHT = "EAF2F8"
BLUE_PALE = "F4F8FB"
GRID = "B8C7D3"
TEXT = "222222"
MUTED = "666666"
AMBER = "8A5A00"
AMBER_FILL = "FFF4D8"
AMBER_BORDER = "D9A441"
GREEN_FILL = "EAF6EF"
GRAY_FILL = "F2F2F2"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border_cell(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def cell_margins(cell, top=95, start=120, bottom=95, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_paragraph_border(paragraph, edges, color=NAVY, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in edges:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    run._r.append(instr)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def paragraph(text="", style=None, size=10.5, bold=False, color=TEXT, align=None, after=5):
    p = DOC.add_paragraph(style=style)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def heading(number, title, level=1):
    text = f"{number}  {title}" if number else title
    p = DOC.add_paragraph(style=f"Heading {min(level, 3)}")
    p.paragraph_format.space_before = Pt(13 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5 if level == 1 else 3)
    r = p.add_run(text)
    set_run_font(r, size=15 if level == 1 else 12.5, bold=True, color=NAVY_DARK if level == 1 else NAVY)
    keep_with_next(p)
    if level == 1:
        add_paragraph_border(p, ("bottom",), color=NAVY, size="8")
    return p


def bullet(text):
    p = DOC.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.36)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.04
    r = p.add_run(text)
    set_run_font(r, size=10.0, color=TEXT)
    return p


def numbered(text):
    p = DOC.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.78)
    p.paragraph_format.first_line_indent = Cm(-0.38)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.04
    r = p.add_run(text)
    set_run_font(r, size=10.0, color=TEXT)
    return p


def make_table(headers, rows, widths, header_fill=BLUE_LIGHT, font_size=9.3):
    table = DOC.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    set_table_geometry(table, widths)
    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            mark_header_row(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade_cell(cell, header_fill if row_index == 0 else "FFFFFF")
            border_cell(cell)
            cell_margins(cell, top=90, start=115, bottom=90, end=115)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    set_run_font(
                        r,
                        size=font_size,
                        bold=(row_index == 0),
                        color=NAVY_DARK if row_index == 0 else TEXT,
                    )
    DOC.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def info_table(rows):
    table = DOC.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "内容"
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    set_table_geometry(table, [2100, 7260])
    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            mark_header_row(row)
        for idx, cell in enumerate(row.cells):
            shade_cell(cell, BLUE_LIGHT if row_index == 0 or idx == 0 else "FFFFFF")
            border_cell(cell)
            cell_margins(cell, top=110, start=135, bottom=110, end=135)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_run_font(
                        r,
                        size=10.0,
                        bold=(row_index == 0 or idx == 0),
                        color=NAVY_DARK if row_index == 0 or idx == 0 else TEXT,
                    )
    DOC.add_paragraph().paragraph_format.space_after = Pt(4)


def note_box(title, lines, fill=AMBER_FILL, border=AMBER_BORDER):
    table = DOC.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, [9360])

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, fill)
    border_cell(cell, border, "8")
    cell_margins(cell, top=150, start=135, bottom=135, end=135)

    title_p = cell.paragraphs[0]
    title_p.paragraph_format.space_after = Pt(5)
    title_p.paragraph_format.line_spacing = 1.05
    r = title_p.add_run(title)
    set_run_font(r, size=10.3, bold=True, color=AMBER if border == AMBER_BORDER else NAVY_DARK)

    for idx, line in enumerate(lines):
        p = cell.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.62)
        p.paragraph_format.first_line_indent = Cm(-0.28)
        p.paragraph_format.space_after = Pt(2 if idx < len(lines) - 1 else 0)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(line)
        set_run_font(r, size=9.6, color=TEXT)

    DOC.add_paragraph().paragraph_format.space_after = Pt(4)


def configure():
    section = DOC.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)

    styles = DOC.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        styles[name].font.bold = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("计算机设计大赛参赛作品说明 | 建筑文化探索游戏操作手册")
    set_run_font(r, size=8.5, color=MUTED)
    add_paragraph_border(p, ("bottom",), color=GRID, size="4")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("第 ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "PAGE")
    r = p.add_run(" 页")
    set_run_font(r, size=8.5, color=MUTED)


def cover():
    p = paragraph("参赛作品操作说明书", size=12, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    p.paragraph_format.space_before = Pt(24)

    p = paragraph("建筑文化探索游戏", size=25, bold=True, color=NAVY_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    p.paragraph_format.space_before = Pt(10)
    p = paragraph("操作手册", size=22, bold=True, color=NAVY_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_paragraph_border(p, ("bottom",), color=NAVY, size="12")

    paragraph("玩家 / 评委阅读版", size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

    info_table(
        [
            ("文档用途", "用于计算机设计大赛现场试玩、评审讲解与操作核对"),
            ("作品类型", "Unity 2D 建筑文化探索与图鉴解锁游戏"),
            ("运行平台", "Windows / Unity 2022.3.62f3c1 项目版本"),
            ("核心目标", "探索地图、收集建筑宝藏、提交构建度、点亮图鉴并解锁建筑文化说明"),
            ("阅读对象", "试玩玩家、评委、演示人员"),
            ("文档版本", "V2.0 优化版"),
        ]
    )

    note_box(
        "文档说明",
        [
            "本文档只描述玩家与评委需要看到的操作和界面功能，不展开代码实现。",
            "当前版本中，主菜单“设置”“退出”等按钮在场景中存在但未绑定业务脚本，手册已按“展示/预留入口”标注。",
        ],
        fill=GREEN_FILL,
        border=NAVY,
    )

    DOC.add_section(WD_SECTION_START.NEW_PAGE)


def build_doc():
    configure()
    cover()

    heading("一", "作品试玩概览", 1)
    paragraph(
        "游戏以建筑文化知识为主题，将地图探索、资源收集、图鉴点亮和建筑详情解锁串联为一个完整试玩闭环。玩家在限定时间内离开基地探索，收集建筑宝藏与专用点亮道具，返回图鉴区域提交资源并逐步解锁赵州桥、福建土楼等建筑条目。"
    )
    make_table(
        ["项目", "说明"],
        [
            ("试玩目标", "完成“探索 -> 拾取 -> 返回 -> 提交 -> 点亮 -> 查看详情”的核心流程。"),
            ("主要挑战", "时间倒计时、敌人追击、生命值消耗、墨水资源消耗、背包容量限制。"),
            ("知识呈现", "通过图鉴小图标说明结构部件，通过建筑详情页呈现建筑背景与文化知识。"),
            ("评审看点", "玩法闭环清晰度、文化内容融入程度、UI 可理解性、交互反馈完整度。"),
        ],
        [2100, 7260],
    )

    heading("二", "按键与基础操作", 1)
    make_table(
        ["编号", "操作", "功能", "触发条件 / 说明"],
        [
            ("K-01", "W / ↑", "向上移动", "使用 Unity Vertical 输入轴，角色向上移动并更新最后朝向。"),
            ("K-02", "S / ↓", "向下移动", "用于探索、靠近宝藏、靠近图鉴和躲避敌人。"),
            ("K-03", "A / ←", "向左移动", "使用 Unity Horizontal 输入轴，角色向左移动并更新最后朝向。"),
            ("K-04", "D / →", "向右移动", "角色四方向移动；同时按横向和纵向时，系统会舍弃横向输入，保留纵向移动。"),
            ("K-05", "鼠标左键", "攻击 / 发射墨球", "每次消耗 5 点墨水；墨球沿角色最后朝向发射；操作类 UI 打开时攻击无效。"),
            ("K-06", "F", "交互", "靠近可交互对象并出现提示后使用，可拾取宝藏、打开图鉴或提交资源。"),
            ("K-07", "鼠标左键点击 UI", "点击按钮、选择条目、翻页", "用于主菜单、图鉴主页、建筑详情、提交窗口、弹窗和失败界面。"),
            ("K-08", "背包格右键长按 1 秒", "丢弃背包单个物品", "仅对有物品的格子生效；丢弃物会生成在角色附近，可再次按 F 拾取。"),
            ("K-09", "Enter / Space", "UI 提交输入", "Unity EventSystem 默认 Submit 输入，可用于当前选中 UI 的提交操作；现场建议优先使用鼠标点击。"),
            ("K-10", "Esc", "UI 取消输入", "Unity EventSystem 默认 Cancel 输入；当前主要流程以界面按钮关闭为准。"),
        ],
        [800, 1700, 2300, 4560],
    )
    note_box(
        "基础操作注意",
        [
            "移动不是八方向移动，斜向组合键不会产生斜向速度。",
            "攻击方向与最后移动方向相关，停下后攻击仍会沿最近一次朝向发射。",
            "图鉴、详情页、提交窗口和弹窗打开时，属于操作类 UI，会阻止攻击，避免误触战斗。",
        ],
    )

    heading("三", "界面区域说明", 1)
    make_table(
        ["编号", "界面 / 区域", "显示内容", "玩家作用"],
        [
            ("UI-01", "主菜单", "开始游戏、设置、退出等入口", "进入试玩；部分入口为当前版本展示/预留按钮。"),
            ("UI-02", "探索场景 HUD", "生命条、墨水条、倒计时、背包栏、交互提示", "帮助玩家判断状态、资源、时间和当前可交互对象。"),
            ("UI-03", "生命条", "玩家当前生命值 / 最大生命值", "被攻击后降低；归零后进入失败界面。"),
            ("UI-04", "墨水 / 武器条", "当前墨水值 / 最大墨水值", "攻击消耗墨水；墨水不足时无法发射墨球。"),
            ("UI-05", "倒计时", "剩余时间，低于 60 秒变红", "离开基地后运行，进入基地区域暂停。"),
            ("UI-06", "背包栏", "6 个固定格子的道具图标", "显示已拾取宝藏和点亮道具；满格后不能继续拾取。"),
            ("UI-07", "交互提示框", "F 图标和交互文本", "提示当前可按 F 执行的动作，如“拾起宝藏”“打开图鉴”。"),
            ("UI-08", "图鉴主页", "建筑大图、进度条、点亮槽位、提交按钮、关闭按钮", "提交资源、点亮结构、查看解锁状态。"),
            ("UI-09", "提交选择窗口", "当前背包物品格", "选择背包物品并提交到指定建筑。"),
            ("UI-10", "建筑详情页", "建筑名称、图片、两页介绍文字", "查看建筑文化背景与结构说明。"),
            ("UI-11", "说明弹窗", "首次拾取或小图标说明文字", "阅读宝藏或结构说明，自动关闭或点击关闭。"),
            ("UI-12", "失败界面", "失败提示语、重新开始、返回主菜单", "死亡后恢复试玩流程。"),
        ],
        [850, 1850, 3000, 3660],
    )

    heading("四", "按钮与控件功能清单", 1)
    paragraph("本节逐项列出玩家/评委在试玩中会接触到的按钮、控件和可点击区域。状态列用于说明当前版本中该按钮是否已有明确业务逻辑。")
    make_table(
        ["编号", "位置", "按钮 / 控件", "状态", "作用"],
        [
            ("B-01", "主菜单", "Start Game / 开始游戏", "已绑定", "点击后调用场景加载器进入 GameScene，开始正式试玩。"),
            ("B-02", "主菜单", "Set up / 设置", "展示/预留", "当前场景中存在按钮对象，但未检索到设置面板或脚本绑定；可作为后续音量、画质、操作说明入口。"),
            ("B-03", "主菜单", "Exit / 退出", "展示/预留", "当前场景中存在按钮对象，但未检索到退出脚本绑定；比赛演示时可说明为预留退出入口。"),
            ("B-04", "主菜单", "HomeButton", "展示元素", "场景对象存在，未发现可点击按钮脚本；更像装饰/导航视觉元素。"),
            ("B-05", "探索场景", "F 交互提示", "已绑定", "靠近可交互物时出现；按 F 执行当前对象的交互逻辑。"),
            ("B-06", "探索场景", "建筑宝藏交互", "已绑定", "按 F 拾取普通宝藏；成功后进入背包，并可能弹出首次拾取说明。"),
            ("B-07", "探索场景", "专用点亮道具交互", "已绑定", "按 F 拾取点亮道具；用于图鉴小图标点亮，不作为普通构建度提交。"),
            ("B-08", "探索场景", "图鉴 / 提交点交互", "已绑定", "按 F 打开图鉴；部分交互点会先上交背包普通宝藏，再打开图鉴。"),
            ("B-09", "背包栏", "背包 1-6 格", "已绑定", "显示物品图标；对有物品格右键长按 1 秒可丢弃单个物品。"),
            ("B-10", "图鉴主页", "AddButton / 提交按钮", "已绑定", "按建筑编号打开对应提交选择窗口；同一建筑再次点击会关闭窗口。"),
            ("B-11", "提交窗口", "Slot 1-6 背包格按钮", "已绑定", "第一次点击选中，第二次点击同一格将该普通宝藏提交给当前建筑。"),
            ("B-12", "提交窗口", "专用点亮道具格", "已绑定禁用", "专用点亮道具不能作为普通宝藏提交，窗口中该格按钮会被禁用。"),
            ("B-13", "图鉴主页", "小图标槽位 Button", "已绑定", "未点亮时点击会消耗一个专用点亮道具并点亮；已点亮时点击显示结构说明。"),
            ("B-14", "图鉴主页", "建筑大图 / 详情入口", "已绑定", "建筑完成解锁后可点击进入详情页；未解锁时不可点击。"),
            ("B-15", "图鉴主页", "CloseButton / 关闭", "已绑定", "关闭图鉴系统，恢复探索 HUD 和玩家移动。"),
            ("B-16", "建筑详情页", "NextPageButton / 下一页", "已绑定", "从详情第一页切换到第二页，查看后续图文说明。"),
            ("B-17", "建筑详情页", "PreviousPageButton / 上一页", "已绑定", "从详情第二页返回第一页。"),
            ("B-18", "建筑详情页", "CloseButton / 关闭", "已绑定", "关闭详情页并返回图鉴主页，不直接退出整个图鉴系统。"),
            ("B-19", "说明弹窗", "ClickCloseButton / 点击关闭", "已绑定", "用于小图标说明弹窗；点击后关闭弹窗并恢复图鉴操作。"),
            ("B-20", "失败界面", "旅程再起", "已绑定", "加载 GameScene，重新开始一局试玩。"),
            ("B-21", "失败界面", "稍事休整", "已绑定", "加载 MainScene，返回主菜单。"),
        ],
        [680, 1300, 1900, 1300, 4180],
        font_size=8.5,
    )

    heading("五", "完整游玩流程", 1)
    for step in [
        "在主菜单点击“开始游戏”，进入 GameScene 探索场景。",
        "观察 HUD：确认生命条、墨水条、倒计时、背包栏和交互提示区域。",
        "离开基地后倒计时开始，使用 WASD 或方向键探索地图。",
        "遇到敌人时用鼠标左键发射墨球。注意墨水消耗和攻击方向。",
        "靠近建筑宝藏或专用点亮道具，出现提示后按 F 拾取。",
        "背包接近满格时返回图鉴/基地区域，避免无法继续拾取。",
        "在图鉴或提交点按 F，打开图鉴主页；若是提交交互点，会先提交背包普通宝藏。",
        "点击某建筑的 AddButton，打开对应提交窗口；选择背包中的普通宝藏并再次点击提交。",
        "使用专用点亮道具点击未点亮的小图标；点亮后再次点击小图标查看结构说明。",
        "当建筑进度满且对应小图标全部点亮后，建筑大图恢复正常显示并开放详情入口。",
        "点击已解锁建筑进入详情页，通过上一页/下一页查看完整建筑说明。",
        "关闭详情页返回图鉴主页，或关闭图鉴回到探索场景继续试玩。",
        "若生命归零，进入失败界面；选择“旅程再起”重开，或“稍事休整”回主菜单。",
    ]:
        numbered(step)

    heading("六", "资源与判定规则", 1)
    make_table(
        ["编号", "资源 / 状态", "来源", "功能", "限制"],
        [
            ("R-01", "普通建筑宝藏", "地图拾取", "增加对应建筑构建度，拾取时可提供临时属性加成。", "提交或丢弃后从背包移除，属性加成随之取消。"),
            ("R-02", "专用点亮道具", "地图拾取", "用于点亮图鉴小图标。", "不能在提交窗口作为普通宝藏提交。"),
            ("R-03", "构建度", "提交普通宝藏", "推动建筑进度条增长。", "每个建筑独立累计，满值后仍需完成小图标点亮。"),
            ("R-04", "小图标点亮状态", "消耗专用点亮道具", "展示建筑结构知识点，参与建筑最终解锁判定。", "每个槽位只能点亮一次。"),
            ("R-05", "建筑最终解锁", "构建度满 + 小图标全亮", "开放建筑详情页。", "任一条件未满足时建筑详情入口不可点击。"),
            ("R-06", "背包容量", "固定 6 格", "承载当前拾取物。", "满格时无法拾取新物品。"),
            ("R-07", "倒计时", "离开基地后运行", "形成探索压力。", "进入基地暂停；低于 60 秒文字变红。"),
            ("R-08", "生命值", "被敌人攻击时减少", "决定玩家是否继续试玩。", "归零进入失败界面。"),
            ("R-09", "墨水值", "初始/加成资源", "驱动远程攻击。", "每次攻击消耗 5 点，低于消耗值无法攻击。"),
        ],
        [760, 1550, 1600, 2900, 2550],
        font_size=8.6,
    )

    heading("七", "注意事项与现场演示建议", 1)
    heading("7.1", "玩家注意事项", 2)
    for item in [
        "看到交互提示后再按 F，离目标过远时交互不会触发。",
        "普通宝藏和专用点亮道具用途不同，提交建筑时不要把两类资源混淆。",
        "背包满时地图物品会保留，但不能继续拾取；应先提交、点亮或丢弃。",
        "打开图鉴、详情页、提交窗口或弹窗时，玩家移动/攻击会被限制，建议站在安全区域操作。",
        "攻击消耗墨水，连续无效攻击会影响后续战斗。",
        "离开基地后倒计时运行，临近 60 秒变红时应优先返回提交或展示已解锁内容。",
    ]:
        bullet(item)

    heading("7.2", "评委试玩路线", 2)
    make_table(
        ["阶段", "建议时长", "操作重点", "观察点"],
        [
            ("入场", "1 分钟", "点击开始游戏，移动、攻击、按 F 交互。", "基础操作是否直观，HUD 是否清晰。"),
            ("探索", "2-3 分钟", "拾取宝藏，观察首次拾取说明和背包变化。", "文化信息是否与道具收集关联。"),
            ("提交", "2 分钟", "返回图鉴，点击 AddButton，选择背包格提交。", "资源流转和建筑进度反馈是否清楚。"),
            ("点亮", "1-2 分钟", "使用专用道具点亮小图标，查看结构说明。", "建筑知识点是否可读、可理解。"),
            ("详情", "1 分钟", "打开已解锁建筑详情页，翻页并返回。", "图文说明、按钮路径和返回逻辑是否完整。"),
            ("失败恢复", "30 秒", "触发或展示失败界面按钮。", "重开与回主菜单路径是否明确。"),
        ],
        [1100, 1000, 3400, 3860],
        font_size=8.8,
    )

    note_box(
        "演示口径建议",
        [
            "先展示玩法闭环，再介绍文化内容；评委更容易理解“为什么要收集、为什么要点亮”。",
            "主菜单设置/退出按钮当前为预留入口，若现场被问到，可说明后续扩展方向是音量、画质、按键说明与退出确认。",
        ],
    )

    DOC.save(OUT)


DOC = Document()


if __name__ == "__main__":
    build_doc()
