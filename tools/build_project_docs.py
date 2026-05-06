from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "文档"
DEPENDENCY_DOC = DOCS_DIR / "依赖说明.docx"
RUN_DOC = DOCS_DIR / "运行说明.docx"

NAVY = "1F4E79"
NAVY_DARK = "123852"
BLUE_LIGHT = "EAF2F8"
BLUE_PALE = "F4F8FB"
GRID = "B8C7D3"
TEXT = "222222"
MUTED = "666666"
AMBER = "8A5A00"
AMBER_FILL = "FFF4D8"
GREEN_FILL = "EAF6EF"
WHITE = "FFFFFF"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig")


def project_version() -> str:
    text = read_text(ROOT / "ProjectSettings" / "ProjectVersion.txt")
    match = re.search(r"m_EditorVersion:\s*(.+)", text)
    return match.group(1).strip() if match else "Unity 2022.3 LTS"


def project_settings() -> dict[str, str]:
    result: dict[str, str] = {}
    text = read_text(ROOT / "ProjectSettings" / "ProjectSettings.asset")
    for key in (
        "companyName",
        "productName",
        "defaultScreenWidth",
        "defaultScreenHeight",
        "fullscreenMode",
        "bundleVersion",
        "apiCompatibilityLevel",
        "activeInputHandler",
    ):
        match = re.search(rf"^\s*{re.escape(key)}:\s*(.+)$", text, re.M)
        if match:
            result[key] = match.group(1).strip()
    return result


def manifest_dependencies() -> dict[str, str]:
    data = json.loads(read_text(ROOT / "Packages" / "manifest.json"))
    return dict(data.get("dependencies", {}))


def package_lock() -> dict[str, dict]:
    data = json.loads(read_text(ROOT / "Packages" / "packages-lock.json"))
    return dict(data.get("dependencies", {}))


def build_scenes() -> list[str]:
    text = read_text(ROOT / "ProjectSettings" / "EditorBuildSettings.asset")
    return re.findall(r"path:\s*(Assets/Scenes/[^ \r\n]+\.unity)", text)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border_cell(cell, color: str = GRID, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def cell_margins(cell, top: int = 95, start: int = 120, bottom: int = 95, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    total = sum(widths)
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_paragraph_border(paragraph, edges: tuple[str, ...], color: str = NAVY, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in edges:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), color)


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    run._r.append(instr)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


class StyledDoc:
    def __init__(self, title: str, subtitle: str, header_title: str) -> None:
        self.doc = Document()
        self.title = title
        self.subtitle = subtitle
        self.header_title = header_title
        self.configure()

    def configure(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.55)
        section.bottom_margin = Cm(1.45)
        section.left_margin = Cm(1.75)
        section.right_margin = Cm(1.75)

        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = RGBColor.from_string(TEXT)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.08
        for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
            styles[name].font.name = "Microsoft YaHei"
            styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            styles[name].font.bold = True

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"计算机设计大赛参赛作品说明 | {self.header_title}")
        set_run_font(r, size=8.5, color=MUTED)
        add_paragraph_border(p, ("bottom",), color=GRID, size="4")

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("第 ")
        set_run_font(r, size=8.5, color=MUTED)
        add_field(p, "PAGE")
        r = p.add_run(" 页")
        set_run_font(r, size=8.5, color=MUTED)

    def paragraph(
        self,
        text: str = "",
        style: str | None = None,
        size: float = 10.5,
        bold: bool = False,
        color: str = TEXT,
        align=None,
        after: float = 5,
    ):
        p = self.doc.add_paragraph(style=style)
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.08
        if text:
            r = p.add_run(text)
            set_run_font(r, size=size, bold=bold, color=color)
        return p

    def heading(self, number: str, title: str, level: int = 1):
        text = f"{number}  {title}" if number else title
        p = self.doc.add_paragraph(style=f"Heading {min(level, 3)}")
        p.paragraph_format.space_before = Pt(13 if level == 1 else 8)
        p.paragraph_format.space_after = Pt(5 if level == 1 else 3)
        r = p.add_run(text)
        set_run_font(r, size=15 if level == 1 else 12.5, bold=True, color=NAVY_DARK if level == 1 else NAVY)
        if level == 1:
            add_paragraph_border(p, ("bottom",), color=NAVY, size="8")
        return p

    def bullet(self, text: str):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = Cm(-0.36)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.04
        r = p.add_run(text)
        set_run_font(r, size=10.0, color=TEXT)
        return p

    def numbered(self, text: str):
        p = self.doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.78)
        p.paragraph_format.first_line_indent = Cm(-0.38)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.04
        r = p.add_run(text)
        set_run_font(r, size=10.0, color=TEXT)
        return p

    def table(self, headers: list[str], rows: list[list[str] | tuple[str, ...]], widths: list[int], font_size: float = 9.2):
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = "Table Grid"
        table.autofit = False
        for i, text in enumerate(headers):
            table.rows[0].cells[i].text = text
        for row in rows:
            cells = table.add_row().cells
            for i, text in enumerate(row):
                cells[i].text = str(text)
        set_table_geometry(table, widths)
        for row_index, row in enumerate(table.rows):
            if row_index == 0:
                mark_header_row(row)
            for cell_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                shade_cell(cell, BLUE_LIGHT if row_index == 0 else WHITE)
                border_cell(cell)
                cell_margins(cell, top=90, start=115, bottom=90, end=115)
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    if cell_index == 0 and row_index > 0 and len(headers) == 2:
                        shade_cell(cell, BLUE_PALE)
                    for r in p.runs:
                        set_run_font(
                            r,
                            size=font_size,
                            bold=(row_index == 0 or (cell_index == 0 and len(headers) == 2)),
                            color=NAVY_DARK if row_index == 0 or (cell_index == 0 and len(headers) == 2) else TEXT,
                        )
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return table

    def note_box(self, title: str, lines: list[str], fill: str = AMBER_FILL, border: str = "D9A441") -> None:
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        set_table_geometry(table, [9360])
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, fill)
        border_cell(cell, border, "8")
        cell_margins(cell, top=150, start=135, bottom=135, end=135)

        title_p = cell.paragraphs[0]
        title_p.paragraph_format.space_after = Pt(5)
        title_p.paragraph_format.line_spacing = 1.05
        r = title_p.add_run(title)
        set_run_font(r, size=10.3, bold=True, color=AMBER if fill == AMBER_FILL else NAVY_DARK)

        for idx, line in enumerate(lines):
            p = cell.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.62)
            p.paragraph_format.first_line_indent = Cm(-0.28)
            p.paragraph_format.space_after = Pt(2 if idx < len(lines) - 1 else 0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(line)
            set_run_font(r, size=9.6, color=TEXT)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def cover(self, info_rows: list[tuple[str, str]], note_lines: list[str]) -> None:
        p = self.paragraph("参赛作品配套说明文档", size=12, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
        p.paragraph_format.space_before = Pt(24)

        self.paragraph("《中国建筑录》", size=25, bold=True, color=NAVY_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        p = self.paragraph(self.title, size=22, bold=True, color=NAVY_DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
        add_paragraph_border(p, ("bottom",), color=NAVY, size="12")
        self.paragraph(self.subtitle, size=12, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

        self.table(["项目", "内容"], info_rows, [2100, 7260], font_size=10.0)
        self.note_box("文档说明", note_lines, fill=GREEN_FILL, border=NAVY)
        self.doc.add_section(WD_SECTION_START.NEW_PAGE)

    def save(self, path: Path) -> None:
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(path)


def dependency_doc() -> None:
    settings = project_settings()
    version = project_version()
    manifest = manifest_dependencies()
    lock = package_lock()
    scenes = build_scenes()
    product = settings.get("productName", "Arcitecture")
    resolution = f"{settings.get('defaultScreenWidth', '1920')} x {settings.get('defaultScreenHeight', '1080')}"

    doc = StyledDoc("依赖说明", "开发 / 评审环境核对版", "《中国建筑录》依赖说明")
    doc.cover(
        [
            ("作品名称", "《中国建筑录》"),
            ("Unity 项目名", product),
            ("Unity Editor", version),
            ("运行定位", "PC 端 Unity 2D 像素风探索游戏"),
            ("推荐分辨率", resolution),
            ("文档用途", "用于计算机设计大赛提交、评审环境复现与项目依赖核对"),
        ],
        [
            "本文档只记录运行和开发所需依赖，不展开源码实现细节。",
            "依赖版本以仓库中的 Packages/manifest.json、Packages/packages-lock.json 和 ProjectSettings 为准。",
        ],
    )

    doc.heading("一", "项目基础环境", 1)
    doc.paragraph("项目基于 Unity 2022.3 LTS 系列制作，当前仓库记录的精确编辑器版本如下。评审或复现时建议优先使用相同版本打开，以减少资源导入、2D 包版本和序列化格式差异。")
    doc.table(
        ["项目", "说明"],
        [
            ("项目根目录", str(ROOT)),
            ("Unity Editor 版本", version),
            ("公司名 / 产品名", f"{settings.get('companyName', 'DefaultCompany')} / {product}"),
            ("项目版本号", settings.get("bundleVersion", "1.0")),
            ("推荐运行平台", "Windows PC；Unity Editor Play Mode 或 Windows Standalone 构建"),
            ("推荐窗口尺寸", resolution),
            ("输入系统", "Unity 旧版 Input Manager（activeInputHandler = 0），使用 Horizontal、Vertical、Mouse0、F 等输入"),
        ],
        [2350, 7010],
    )

    doc.heading("二", "Unity Package 依赖", 1)
    doc.paragraph("下表列出 manifest.json 中直接声明的核心包。Unity 打开项目时会根据 manifest 与 lock 文件从 Unity Package Registry 恢复这些依赖。")
    package_rows = [
        ("com.unity.feature.2d", manifest.get("com.unity.feature.2d", ""), "2D 项目特性集合，包含 Sprite、Tilemap、Pixel Perfect、PSD Importer 等 2D 工作流"),
        ("com.unity.textmeshpro", manifest.get("com.unity.textmeshpro", ""), "中文 UI 文本、说明弹窗、按钮标签和 HUD 文本渲染"),
        ("com.unity.ugui", manifest.get("com.unity.ugui", ""), "主菜单、HUD、背包、图鉴、提交窗口、失败界面等 UI 系统"),
        ("com.unity.timeline", manifest.get("com.unity.timeline", ""), "时间轴能力预留，可配合动画和展示流程使用"),
        ("com.unity.visualscripting", manifest.get("com.unity.visualscripting", ""), "Unity 可视化脚本包，随项目依赖保留"),
        ("com.unity.test-framework", manifest.get("com.unity.test-framework", ""), "Unity 测试框架，支持后续编辑器/运行时测试扩展"),
        ("com.unity.ide.visualstudio", manifest.get("com.unity.ide.visualstudio", ""), "Visual Studio 集成与 C# 脚本编辑支持"),
        ("com.unity.ide.rider", manifest.get("com.unity.ide.rider", ""), "JetBrains Rider 集成与 C# 脚本编辑支持"),
        ("com.unity.collab-proxy", manifest.get("com.unity.collab-proxy", ""), "Unity 协作/版本控制相关集成包"),
    ]
    doc.table(["包名", "版本", "用途"], package_rows, [2550, 1100, 5710], font_size=8.8)

    doc.heading("三", "2D 与内置模块依赖", 1)
    doc.paragraph("项目大量使用 Unity 内置模块完成 2D 场景、碰撞、动画、UI 和资源读写。下表按功能归纳，不要求单独安装。")
    doc.table(
        ["模块类别", "涉及包 / 模块", "用途说明"],
        [
            ("2D 地图与精灵", "2D Sprite、2D Tilemap、2D Animation、2D PSD Importer、2D Pixel Perfect", "地图绘制、角色与敌人动画、建筑资源、像素风画面表现"),
            ("物理与交互", "com.unity.modules.physics2d、Collider2D、Rigidbody2D、Trigger", "玩家移动、拾取触发、敌人追踪、交互距离判断和碰撞逻辑"),
            ("动画系统", "com.unity.modules.animation、Animator、AnimationClip", "角色行走、攻击、敌人动作、主菜单淡入淡出和场景过渡"),
            ("UI 与文本", "com.unity.modules.ui、UGUI、TextMeshPro", "按钮、背包格、进度条、倒计时、图鉴、说明弹窗和失败界面"),
            ("数据与资源", "JsonSerialize、ImageConversion、ScreenCapture、UnityWebRequest 模块", "本地数据序列化、图片资源处理、截图/相册功能和资源扩展能力"),
            ("音频与表现", "Audio、ParticleSystem、Video 模块", "音效、特效和后续展示内容扩展"),
        ],
        [1750, 2900, 4710],
        font_size=8.6,
    )

    doc.heading("四", "资源目录依赖", 1)
    doc.table(
        ["目录 / 文件", "说明"],
        [
            ("Assets/Scenes", "包含 MainScene、GameScene、DeadScene 三个构建场景，是运行项目的核心入口。"),
            ("Assets/Scripts", "包含主菜单控制、玩家移动/攻击/交互、背包、图鉴、敌人、倒计时和 UI 管理脚本。"),
            ("Assets/File", "主要美术与游戏资源目录，包含 UIResources、Prefab、Prop、TileMap、TexturePackage 等素材。"),
            ("Assets/TextMesh Pro", "TextMeshPro 字体、材质、Sprite Asset、Shader 和示例资源；中文 UI 依赖 TMP 组件显示。"),
            ("Packages/manifest.json", "直接依赖声明文件，Unity Package Manager 首先读取该文件恢复依赖。"),
            ("Packages/packages-lock.json", "依赖锁定文件，记录直接和间接包版本，例如 2D Animation 9.2.0、Tilemap Extras 3.1.3。"),
            ("ProjectSettings", "记录 Unity 版本、构建场景、输入管理、画面尺寸、质量设置和项目全局配置。"),
        ],
        [2700, 6660],
    )

    doc.heading("五", "场景与构建依赖", 1)
    scene_rows = []
    for index, scene in enumerate(scenes):
        scene_name = Path(scene).stem
        if scene_name == "MainScene":
            usage = "主菜单场景，包含开始游戏入口和主场景视觉资源。"
        elif scene_name == "GameScene":
            usage = "正式试玩场景，承载探索、战斗、拾取、图鉴和 HUD。"
        elif scene_name == "DeadScene":
            usage = "失败结算场景，提供重新开始和返回主菜单按钮。"
        else:
            usage = "项目构建场景。"
        scene_rows.append((str(index), scene, usage))
    doc.table(["构建序号", "场景路径", "用途"], scene_rows, [900, 3000, 5460], font_size=8.8)
    doc.note_box(
        "构建注意",
        [
            "Build Settings 中应保持 MainScene 为 0，GameScene 为 1，DeadScene 为 2；主菜单和失败界面的场景跳转依赖该顺序或对应场景名。",
            "若打开项目后场景列表为空，应在 Unity 中依次添加 Assets/Scenes/MainScene.unity、GameScene.unity、DeadScene.unity。",
        ],
    )

    doc.heading("六", "依赖恢复与检查步骤", 1)
    for step in [
        "使用 Unity Hub 选择 Add/Open project，打开项目根目录 D:\\Workspace\\Arcitecture。",
        f"确认 Unity Editor 版本为 {version}；若本机没有该版本，优先安装同一 LTS 补丁版本。",
        "首次打开时等待 Unity 导入 Assets 并通过 Package Manager 恢复 manifest.json 中的依赖。",
        "打开 Package Manager，确认 TextMeshPro、2D Feature、UGUI、Visual Studio/Rider 支持包无红色错误。",
        "打开 Build Settings，确认三个场景按 MainScene、GameScene、DeadScene 顺序加入并勾选。",
        "打开 MainScene 后进入 Play Mode，若能显示主菜单并点击开始游戏进入 GameScene，则依赖恢复完成。",
    ]:
        doc.numbered(step)

    doc.heading("七", "常见依赖问题", 1)
    doc.table(
        ["问题现象", "可能原因", "处理建议"],
        [
            ("Package Manager 报错或包未下载", "网络或 Unity Registry 源不可用", "检查网络与 Unity Hub 登录状态，必要时重开项目让 Package Manager 重新解析。"),
            ("TMP 文本显示异常", "TextMeshPro 资源未导入或字体资源缺失", "确认 Assets/TextMesh Pro 目录存在，并在 Unity 中执行 TMP Essential Resources 导入检查。"),
            ("场景跳转失败", "Build Settings 场景未加入或顺序错误", "按 MainScene、GameScene、DeadScene 顺序加入构建列表。"),
            ("移动或攻击无响应", "Input Manager 配置异常或当前 UI 阻塞玩法输入", "确认 activeInputHandler 为旧版输入，并关闭图鉴、提交窗口、详情页等阻塞 UI 后再操作。"),
            ("像素资源或 Prefab 丢失", "Assets/File 目录缺失或资源 GUID 被破坏", "确认完整拉取仓库资源和 .meta 文件，不要单独移动图片/Prefab 文件。"),
        ],
        [2200, 2500, 4660],
        font_size=8.6,
    )

    doc.save(DEPENDENCY_DOC)


def run_doc() -> None:
    settings = project_settings()
    version = project_version()
    scenes = build_scenes()
    resolution = f"{settings.get('defaultScreenWidth', '1920')} x {settings.get('defaultScreenHeight', '1080')}"

    doc = StyledDoc("运行说明", "玩家 / 评委试玩版", "《中国建筑录》运行说明")
    doc.cover(
        [
            ("作品名称", "《中国建筑录》"),
            ("Unity 项目名", settings.get("productName", "Arcitecture")),
            ("运行平台", "Windows PC / Unity Editor Play Mode"),
            ("Unity Editor", version),
            ("推荐分辨率", resolution),
            ("文档用途", "用于说明项目如何启动、试玩路径和主要操作方式"),
        ],
        [
            "本文档面向比赛评委、试玩玩家和现场演示人员，重点说明如何运行游戏与完成核心试玩闭环。",
            "运行路径以当前仓库场景和脚本为准：从 MainScene 进入 GameScene，失败后进入 DeadScene。",
        ],
    )

    doc.heading("一", "运行环境准备", 1)
    doc.paragraph("运行前应保证 Unity 项目依赖已经恢复，且构建场景顺序正确。项目本体不需要额外数据库或服务器，主要依赖 Unity 本地资源、场景和脚本运行。")
    doc.table(
        ["项目", "说明"],
        [
            ("Unity Editor", version),
            ("项目目录", str(ROOT)),
            ("推荐平台", "Windows PC"),
            ("推荐分辨率", resolution),
            ("输入方式", "键盘 + 鼠标；使用旧版 Input Manager"),
            ("运行入口", "Assets/Scenes/MainScene.unity"),
        ],
        [2300, 7060],
    )

    doc.heading("二", "Unity Editor 运行步骤", 1)
    for step in [
        "打开 Unity Hub，选择项目根目录 D:\\Workspace\\Arcitecture。",
        f"使用 Unity {version} 或兼容的 2022.3 LTS 版本打开项目，等待资源导入和包依赖恢复完成。",
        "在 Project 面板打开 Assets/Scenes/MainScene.unity。",
        "进入 File > Build Settings，确认场景列表依次为 MainScene、GameScene、DeadScene 且均已勾选。",
        "点击 Unity 顶部 Play 按钮进入主菜单。",
        "在主菜单点击“Start Game / 开始游戏”按钮，进入 GameScene 开始试玩。",
    ]:
        doc.numbered(step)

    doc.heading("三", "场景说明", 1)
    scene_rows = []
    for index, scene in enumerate(scenes):
        name = Path(scene).stem
        if name == "MainScene":
            desc = "主菜单场景。玩家从此处开始游戏；当前版本设置/退出按钮作为展示或预留入口。"
        elif name == "GameScene":
            desc = "核心试玩场景。包含探索地图、HUD、玩家移动攻击、敌人、拾取、背包、图鉴、提交和建筑详情。"
        elif name == "DeadScene":
            desc = "失败界面。生命归零后进入，可选择“旅程再起”重新开始或“稍事休整”返回主菜单。"
        else:
            desc = "项目构建场景。"
        scene_rows.append((str(index), name, scene, desc))
    doc.table(["序号", "场景", "路径", "说明"], scene_rows, [650, 1350, 2800, 4560], font_size=8.5)

    doc.heading("四", "基础操作", 1)
    doc.table(
        ["操作", "按键 / 鼠标", "功能说明"],
        [
            ("移动", "W / A / S / D 或方向键", "控制角色四方向移动；同时按横向和纵向时优先保留纵向移动。"),
            ("攻击", "鼠标左键", "发射墨球攻击敌人，每次消耗 5 点墨水；图鉴、详情页、提交窗口等玩法阻塞 UI 打开时无法攻击。"),
            ("交互", "F", "靠近可交互对象并出现提示后触发拾取、打开图鉴或提交资源等行为。"),
            ("UI 点击", "鼠标左键", "点击主菜单、图鉴、提交窗口、详情页、失败界面中的按钮或格子。"),
            ("背包丢弃", "背包格右键长按 1 秒", "对已有物品的背包格生效，丢弃后物品生成在角色附近，可再次靠近按 F 拾取。"),
            ("详情翻页", "鼠标点击上一页 / 下一页", "在建筑详情页切换说明页面，查看完整建筑文化内容。"),
        ],
        [1650, 2200, 5510],
        font_size=8.7,
    )

    doc.heading("五", "完整试玩流程", 1)
    for step in [
        "在主菜单点击“开始游戏”，进入 GameScene 探索场景。",
        "观察 HUD 中的生命条、墨水条、倒计时、背包栏和交互提示区域。",
        "离开基地后开始探索，使用 WASD 或方向键移动，寻找建筑宝藏和专用点亮道具。",
        "遇到敌人时使用鼠标左键发射墨球，注意墨水消耗和角色最后朝向。",
        "靠近宝藏、点亮道具或图鉴交互点，出现 F 提示后按 F 触发交互。",
        "背包装入普通建筑宝藏或专用点亮道具；若背包接近满格，应返回图鉴/基地区域处理资源。",
        "在图鉴或提交点按 F 打开图鉴主页，点击建筑对应的 AddButton 打开提交窗口。",
        "在提交窗口选择普通宝藏并再次点击对应格子，将资源提交给当前建筑，推动构建进度。",
        "使用专用点亮道具点击未点亮的小图标；点亮后再次点击小图标可查看结构说明。",
        "当建筑进度满且对应小图标全部点亮后，建筑大图恢复正常显示，并开放详情入口。",
        "点击已解锁建筑进入详情页，通过上一页/下一页查看建筑历史背景和结构知识。",
        "关闭详情页返回图鉴主页，或关闭图鉴回到探索场景继续收集与战斗。",
        "若生命归零，进入 DeadScene；点击“旅程再起”重开 GameScene，或点击“稍事休整”返回 MainScene。",
    ]:
        doc.numbered(step)

    doc.heading("六", "核心界面与按钮", 1)
    doc.table(
        ["界面 / 区域", "主要内容", "玩家作用"],
        [
            ("主菜单", "开始游戏、设置、退出等入口", "点击开始游戏进入正式试玩；设置/退出为当前版本预留或展示入口。"),
            ("探索 HUD", "生命、墨水、倒计时、背包、交互提示", "判断生存状态、资源消耗、剩余时间和当前可交互对象。"),
            ("背包栏", "6 个固定物品格", "显示已拾取资源；满格后无法继续拾取新物品，可提交、点亮或右键长按丢弃。"),
            ("图鉴主页", "建筑大图、进度条、结构槽位、提交按钮", "提交资源、点亮结构节点、查看建筑解锁状态。"),
            ("提交窗口", "当前背包物品格", "选择普通建筑宝藏提交给指定建筑；专用点亮道具不作为普通宝藏提交。"),
            ("建筑详情页", "建筑图片、说明文字、翻页和关闭按钮", "阅读已解锁建筑的文化背景、结构逻辑和图文说明。"),
            ("失败界面", "重新开始、返回主菜单", "生命归零后的恢复入口。"),
        ],
        [1800, 3000, 4560],
        font_size=8.8,
    )

    doc.heading("七", "现场演示建议", 1)
    doc.note_box(
        "推荐演示路线",
        [
            "先展示主菜单进入 GameScene，再快速说明 HUD 中生命、墨水、倒计时和背包的意义。",
            "用一次移动、一次攻击、一次 F 拾取建立基础操作认知。",
            "返回图鉴区域提交普通宝藏，再使用专用道具点亮小图标，强调“收集 - 修复 - 解锁知识”的闭环。",
            "最后打开已解锁建筑详情页，展示文化内容如何从玩法结果自然出现。",
        ],
        fill=GREEN_FILL,
        border=NAVY,
    )
    doc.table(
        ["阶段", "建议时长", "操作重点", "展示价值"],
        [
            ("启动", "30 秒", "打开 MainScene 并点击开始游戏", "说明项目可运行、入口清晰。"),
            ("探索", "1-2 分钟", "移动、攻击、拾取、观察 HUD", "展示搜打撤基础体验和资源压力。"),
            ("提交", "1 分钟", "打开图鉴并提交背包宝藏", "展示建筑修复进度反馈。"),
            ("点亮", "1 分钟", "使用专用道具点亮结构小图标", "展示建筑结构知识点与玩法绑定。"),
            ("详情", "30 秒", "打开建筑详情并翻页", "展示文化内容沉淀。"),
            ("失败恢复", "30 秒", "展示重新开始或返回主菜单", "说明闭环完整。"),
        ],
        [1100, 1000, 3200, 4060],
        font_size=8.6,
    )

    doc.heading("八", "常见运行问题", 1)
    doc.table(
        ["问题现象", "可能原因", "解决方式"],
        [
            ("打开项目后控制台出现包错误", "Unity 包依赖未恢复完成", "等待 Package Manager 完成解析，必要时重启 Unity 重新导入。"),
            ("点击开始游戏没有进入探索场景", "SceneLoader 或 Build Settings 场景序号异常", "确认 MainScene 为 0、GameScene 为 1，且 Start Game 按钮绑定脚本存在。"),
            ("TMP 文本缺字或显示异常", "TextMeshPro 资源未正确导入", "检查 Assets/TextMesh Pro 目录，必要时导入 TMP Essential Resources。"),
            ("角色不能攻击", "墨水不足或操作类 UI 正在打开", "关闭图鉴/详情/提交窗口等 UI，或确认墨水值不低于单次消耗。"),
            ("按 F 没有反应", "距离交互对象过远或未进入触发区域", "靠近目标，看到 F 图标和提示文本后再按 F。"),
            ("背包无法继续拾取", "6 格背包已满", "返回图鉴提交资源、点亮结构槽位，或对格子右键长按 1 秒丢弃。"),
            ("失败后无法回到游戏", "DeadScene 按钮或场景名异常", "确认 GameScene 和 MainScene 场景文件存在且已加入 Build Settings。"),
        ],
        [2200, 2500, 4660],
        font_size=8.6,
    )

    doc.save(RUN_DOC)


def main() -> None:
    dependency_doc()
    run_doc()
    print(f"Generated: {DEPENDENCY_DOC}")
    print(f"Generated: {RUN_DOC}")


if __name__ == "__main__":
    main()
